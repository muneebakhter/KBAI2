#!/usr/bin/env python3
"""
Document processing utilities for DARKBO
Handles PDF and DOCX document text extraction and cleaning
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Optional, Tuple, Dict
import tempfile

# Document processing imports with fallbacks
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes documents and extracts clean text"""
    
    def __init__(self):
        self.chunk_size = 1000  # Target chunk size in characters
        self.chunk_overlap = 100  # Overlap between chunks
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported"""
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            return HAS_PDF
        elif ext in ['.docx', '.doc']:
            return HAS_DOCX
        return False
    
    def extract_text(self, file_path: str) -> Tuple[str, List[str]]:
        """
        Extract text from document file
        Returns: (full_text, list_of_chunks)
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            return self._extract_pdf_text(file_path)
        elif ext in ['.docx', '.doc']:
            return self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _extract_pdf_text(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text from PDF file"""
        if not HAS_PDF:
            raise ImportError("PyPDF2 not available for PDF processing")
        
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        continue
        
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise ValueError(f"Failed to process PDF: {e}")
        
        # Join all text and clean it
        full_text = "\n\n".join(text_content)
        cleaned_text = self._clean_text(full_text)
        
        # Create chunks
        chunks = self._create_chunks(cleaned_text)
        
        return cleaned_text, chunks
    
    def _extract_docx_text(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text from DOCX file"""
        if not HAS_DOCX:
            raise ImportError("python-docx not available for DOCX processing")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
        
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise ValueError(f"Failed to process DOCX: {e}")
        
        # Join all text and clean it
        full_text = "\n\n".join(text_content)
        cleaned_text = self._clean_text(full_text)
        
        # Create chunks
        chunks = self._create_chunks(cleaned_text)
        
        return cleaned_text, chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\"\'\/]', ' ', text)
        
        # Remove multiple dots/dashes
        text = re.sub(r'\.{3,}', '...', text)
        text = re.sub(r'-{3,}', '---', text)
        
        # Clean up spacing around punctuation
        text = re.sub(r'\s+([,.;:!?])', r'\1', text)
        text = re.sub(r'([,.;:!?])\s*([,.;:!?])', r'\1 \2', text)
        
        # Remove lines that are too short (likely formatting artifacts)
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if len(line) > 10 or (len(line) > 3 and any(c.isalpha() for c in line)):
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines).strip()
    
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []
        
        chunks = []
        text_length = len(text)
        
        # If text is shorter than chunk size, return as single chunk
        if text_length <= self.chunk_size:
            return [text]
        
        start = 0
        while start < text_length:
            end = start + self.chunk_size
            
            # If this is not the last chunk, try to end at a sentence boundary
            if end < text_length:
                # Look for sentence ending in the last 200 characters of the chunk
                chunk_text = text[start:end]
                sentence_endings = [m.end() for m in re.finditer(r'[.!?]\s+', chunk_text[-200:])]
                
                if sentence_endings:
                    # Use the last sentence ending found
                    last_sentence_end = sentence_endings[-1]
                    end = start + len(chunk_text) - 200 + last_sentence_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start forward, accounting for overlap
            start = max(start + self.chunk_size - self.chunk_overlap, end)
            
            # Prevent infinite loop
            if start >= text_length:
                break
        
        return chunks
    
    def get_document_metadata(self, file_path: str) -> Dict[str, any]:
        """Extract document metadata"""
        path = Path(file_path)
        
        metadata = {
            "filename": path.name,
            "file_size": path.stat().st_size,
            "file_extension": path.suffix.lower(),
            "supported": self.is_supported(path.name)
        }
        
        # Add format-specific metadata
        ext = path.suffix.lower()
        if ext == '.pdf' and HAS_PDF:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        "page_count": len(pdf_reader.pages),
                        "format": "PDF"
                    })
            except:
                pass
        
        elif ext in ['.docx', '.doc'] and HAS_DOCX:
            try:
                doc = DocxDocument(file_path)
                metadata.update({
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "format": "DOCX"
                })
            except:
                pass
        
        return metadata


def process_document_for_kb(file_path: str, article_title: str = None) -> Tuple[str, List[str], Dict]:
    """
    High-level function to process a document for knowledge base ingestion
    Returns: (cleaned_text, chunks, metadata)
    """
    processor = DocumentProcessor()
    
    # Generate article title if not provided
    if not article_title:
        article_title = Path(file_path).stem.replace('_', ' ').replace('-', ' ').title()
    
    # Extract text and create chunks
    full_text, chunks = processor.extract_text(file_path)
    
    # Get metadata
    metadata = processor.get_document_metadata(file_path)
    metadata['article_title'] = article_title
    metadata['chunk_count'] = len(chunks)
    metadata['text_length'] = len(full_text)
    
    return full_text, chunks, metadata